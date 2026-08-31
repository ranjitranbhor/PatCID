#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Turn user-supplied documents into page images (and text) for the pipeline.

Supported inputs
----------------
``.pdf``           rendered page by page with PyMuPDF at a configurable DPI
``.docx`` / ``.doc``  converted to PDF with LibreOffice when available, so page
                   numbers and layout are preserved; otherwise the embedded
                   images are extracted and each is treated as one "page"
``.png``/``.jpg``/... a single-page image document

Rendering DPI defaults to 300, matching DECIMER-Segmentation's own PDF handling.
"""

from __future__ import annotations

import hashlib
import shutil
import subprocess
import tempfile
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, Iterator, List, Optional

import numpy as np

PDF_SUFFIXES = {".pdf"}
WORD_SUFFIXES = {".docx", ".doc", ".docm", ".odt", ".rtf"}
IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".gif", ".webp"}
SUPPORTED_SUFFIXES = PDF_SUFFIXES | WORD_SUFFIXES | IMAGE_SUFFIXES

DEFAULT_DPI = 300


@dataclass
class Page:
    """One rendered page of a document."""

    number: int  # 1-based
    image: np.ndarray  # RGB, uint8, shape (h, w, 3)
    dpi: int
    text: str = ""
    source: str = "render"  # "render" | "embedded-image"

    @property
    def height(self) -> int:
        return int(self.image.shape[0])

    @property
    def width(self) -> int:
        return int(self.image.shape[1])


@dataclass
class Document:
    """A user-supplied document, resolved to a page sequence on demand."""

    path: Path
    dpi: int = DEFAULT_DPI
    _temp_dir: Optional[Path] = field(default=None, repr=False)

    @property
    def name(self) -> str:
        return self.path.name

    @property
    def suffix(self) -> str:
        return self.path.suffix.lower()

    def content_hash(self) -> str:
        """SHA-256 of the file contents - the cache key for this document."""
        digest = hashlib.sha256()
        with open(self.path, "rb") as handle:
            for chunk in iter(lambda: handle.read(1 << 20), b""):
                digest.update(chunk)
        return digest.hexdigest()

    # -- page production -------------------------------------------------------

    def pages(self, page_numbers: Optional[List[int]] = None) -> Iterator[Page]:
        """Yield pages, optionally restricted to a 1-based page selection."""
        if self.suffix in PDF_SUFFIXES:
            yield from _pdf_pages(self.path, self.dpi, page_numbers)
        elif self.suffix in WORD_SUFFIXES:
            yield from self._word_pages(page_numbers)
        elif self.suffix in IMAGE_SUFFIXES:
            if page_numbers and 1 not in page_numbers:
                return
            yield _image_page(self.path, self.dpi)
        else:
            raise ValueError(
                f"Unsupported document type {self.suffix!r} for {self.path}. "
                f"Supported: {sorted(SUPPORTED_SUFFIXES)}"
            )

    def page_count(self) -> int:
        if self.suffix in PDF_SUFFIXES:
            import pymupdf

            with pymupdf.open(self.path) as pdf:
                return pdf.page_count
        if self.suffix in IMAGE_SUFFIXES:
            return 1
        converted = self._converted_pdf()
        if converted is not None:
            import pymupdf

            with pymupdf.open(converted) as pdf:
                return pdf.page_count
        return len(_docx_embedded_images(self.path))

    def _word_pages(self, page_numbers: Optional[List[int]]) -> Iterator[Page]:
        converted = self._converted_pdf()
        if converted is not None:
            yield from _pdf_pages(converted, self.dpi, page_numbers)
            return
        # No LibreOffice: fall back to the images embedded in the .docx archive.
        # Page numbers are then image indices, and no page-level text is available.
        for index, image in enumerate(_docx_embedded_images(self.path), start=1):
            if page_numbers and index not in page_numbers:
                continue
            yield Page(
                number=index,
                image=image,
                dpi=self.dpi,
                text="",
                source="embedded-image",
            )

    def _converted_pdf(self) -> Optional[Path]:
        """Convert a Word document to PDF via LibreOffice, if it is installed."""
        if self._temp_dir is not None:
            candidate = self._temp_dir / (self.path.stem + ".pdf")
            return candidate if candidate.exists() else None
        soffice = shutil.which("soffice") or shutil.which("libreoffice")
        if soffice is None:
            return None
        self._temp_dir = Path(tempfile.mkdtemp(prefix="structure_finder_"))
        try:
            subprocess.run(
                [
                    soffice,
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    str(self._temp_dir),
                    str(self.path),
                ],
                check=True,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
                timeout=600,
            )
        except (subprocess.CalledProcessError, subprocess.TimeoutExpired, OSError):
            return None
        candidate = self._temp_dir / (self.path.stem + ".pdf")
        return candidate if candidate.exists() else None

    def cleanup(self) -> None:
        if self._temp_dir is not None and self._temp_dir.exists():
            shutil.rmtree(self._temp_dir, ignore_errors=True)
        self._temp_dir = None

    # -- full text -------------------------------------------------------------

    def full_text(self) -> str:
        """Extract the document text (used for the SMILES/InChI text scan)."""
        if self.suffix in PDF_SUFFIXES:
            return _pdf_text(self.path)
        if self.suffix in WORD_SUFFIXES:
            converted = self._converted_pdf()
            if converted is not None:
                return _pdf_text(converted)
            return _docx_text(self.path)
        return ""

    def metadata(self) -> Dict[str, object]:
        return {
            "filename": self.name,
            "path": str(self.path),
            "sha256": self.content_hash(),
            "bytes": self.path.stat().st_size,
            "dpi": self.dpi,
        }


# ---------------------------------------------------------------------------
# Backend helpers
# ---------------------------------------------------------------------------


def _pdf_pages(
    path: Path, dpi: int, page_numbers: Optional[List[int]] = None
) -> Iterator[Page]:
    import pymupdf

    wanted = set(page_numbers) if page_numbers else None
    with pymupdf.open(path) as pdf:
        matrix = pymupdf.Matrix(dpi / 72, dpi / 72)
        for index in range(pdf.page_count):
            number = index + 1
            if wanted is not None and number not in wanted:
                continue
            page = pdf[index]
            pixmap = page.get_pixmap(matrix=matrix, alpha=False)
            image = np.frombuffer(pixmap.samples, dtype=np.uint8).reshape(
                pixmap.h, pixmap.w, pixmap.n
            )
            if pixmap.n == 1:
                image = np.repeat(image, 3, axis=2)
            elif pixmap.n == 4:
                image = image[:, :, :3]
            yield Page(
                number=number,
                image=np.ascontiguousarray(image),
                dpi=dpi,
                text=page.get_text() or "",
                source="render",
            )


def _pdf_text(path: Path) -> str:
    import pymupdf

    with pymupdf.open(path) as pdf:
        return "\n".join(page.get_text() or "" for page in pdf)


def _image_page(path: Path, dpi: int) -> Page:
    from PIL import Image

    with Image.open(path) as handle:
        image = np.array(handle.convert("RGB"))
    return Page(number=1, image=image, dpi=dpi, text="", source="render")


def _docx_embedded_images(path: Path) -> List[np.ndarray]:
    """Extract raster images stored inside a .docx/.odt archive."""
    import io

    from PIL import Image

    images: List[np.ndarray] = []
    try:
        archive = zipfile.ZipFile(path)
    except (zipfile.BadZipFile, OSError):
        return images
    with archive:
        media = sorted(
            name
            for name in archive.namelist()
            if (name.startswith("word/media/") or name.startswith("Pictures/"))
            and Path(name).suffix.lower() in IMAGE_SUFFIXES
        )
        for name in media:
            try:
                with Image.open(io.BytesIO(archive.read(name))) as handle:
                    images.append(np.array(handle.convert("RGB")))
            except Exception:
                continue
    return images


def _docx_text(path: Path) -> str:
    try:
        import docx  # python-docx
    except ImportError:
        return ""
    try:
        document = docx.Document(str(path))
    except Exception:
        return ""
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def collect_documents(
    inputs, dpi: int = DEFAULT_DPI, recursive: bool = True
) -> List[Document]:
    """Expand files and directories into a list of :class:`Document`.

    Args:
        inputs: A path, or an iterable of paths, to files and/or directories.
            A single path may be given as a bare string or ``Path`` - iterating
            a string yields characters, so accepting one explicitly avoids a
            baffling "No such file or directory: c" further down.
    """
    if isinstance(inputs, (str, Path)):
        inputs = [inputs]
    else:
        inputs = list(inputs)

    documents: List[Document] = []
    seen: set = set()
    for item in inputs:
        if not isinstance(item, (str, Path)):
            raise TypeError(
                f"Each input must be a path, got {type(item).__name__}: {item!r}"
            )
        path = Path(item).expanduser()
        if path.is_dir():
            pattern = "**/*" if recursive else "*"
            candidates = sorted(
                child
                for child in path.glob(pattern)
                if child.is_file() and child.suffix.lower() in SUPPORTED_SUFFIXES
            )
        elif path.is_file():
            candidates = [path]
        else:
            raise FileNotFoundError(f"No such file or directory: {path}")
        for candidate in candidates:
            resolved = candidate.resolve()
            if resolved in seen:
                continue
            if resolved.suffix.lower() not in SUPPORTED_SUFFIXES:
                continue
            seen.add(resolved)
            documents.append(Document(path=resolved, dpi=dpi))
    return documents
